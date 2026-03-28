# -*- coding: utf-8 -*-
"""
tao_bo_de_final.py
─────────────────────────────────────────────────────────────────
Bộ đề HSG Ngữ Văn 7 – 2025-2026
  Trang 1 : Đề thi (thơ 2 cột, vừa 1 trang A4)
  Trang 2+: Hướng dẫn chấm (bảng Phần | Nội dung cần đạt | Điểm)

Usage:
  python tao_bo_de_final.py           # tạo cả 3 đề
  python tao_bo_de_final.py --de 1    # chỉ tạo Đề 1
  python tao_bo_de_final.py --de 2 3  # tạo Đề 2 và Đề 3
  python tao_bo_de_final.py --validate # chỉ kiểm tra điểm, không tạo file
─────────────────────────────────────────────────────────────────
"""
import argparse
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from parse_poem import parse_poem, split_two_columns

FONT = 'Times New Roman'
SZ   = 11

# Thư mục xuất file .doc
OUTPUT_DIR = Path(__file__).parent / 'output'

# Đọc bài thơ từ file txt
_POEM_PATH = Path(__file__).parent / 'nho_ba.txt'
_poem      = parse_poem(str(_POEM_PATH))
POEM_TITLE = _poem['title']        # 'NHỚ BÀ'
POEM_SOURCE = _poem['source']      # '(Nhớ bà – Trương Anh Tú ...)'
LEFT_STANZAS, RIGHT_STANZAS = split_two_columns(_poem['stanzas'], left_count=3)

# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def setup(doc):
    s = doc.styles['Normal']
    s.font.name = FONT; s.font.size = Pt(SZ)
    s.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(3.0); sec.right_margin = Cm(2.0)


def R(para, text, bold=False, italic=False, underline=False, sz=None):
    run = para.add_run(text)
    run.bold = bold; run.italic = italic; run.underline = underline
    run.font.name = FONT; run.font.size = Pt(sz or SZ)
    run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), FONT)
    return run


def P(doc, align=None, sb=0, sa=0, fi=None):
    para = doc.add_paragraph()
    if align: para.alignment = align
    para.paragraph_format.space_before      = Pt(sb)
    para.paragraph_format.space_after       = Pt(sa)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if fi is not None:
        para.paragraph_format.first_line_indent = Cm(fi)
    return para


def no_border(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    bd = OxmlElement('w:tblBorders')
    for e in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{e}')
        el.set(qn('w:val'),'none'); el.set(qn('w:sz'),'0')
        el.set(qn('w:space'),'0'); el.set(qn('w:color'),'auto')
        bd.append(el)
    tblPr.append(bd)


def full_border(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    bd = OxmlElement('w:tblBorders')
    for e in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{e}')
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4')
        el.set(qn('w:space'),'0'); el.set(qn('w:color'),'000000')
        bd.append(el)
    tblPr.append(bd)


def clr(cell):
    for para in list(cell.paragraphs):
        para._element.getparent().remove(para._element)


def CP(cell, align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=0):
    para = cell.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before      = Pt(sb)
    para.paragraph_format.space_after       = Pt(sa)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return para


def hrule(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(2)
    bdr = OxmlElement('w:pBdr')
    btm = OxmlElement('w:bottom')
    btm.set(qn('w:val'),'single'); btm.set(qn('w:sz'),'6')
    btm.set(qn('w:space'),'1');   btm.set(qn('w:color'),'000000')
    bdr.append(btm)
    para._p.get_or_add_pPr().append(bdr)


def pagebreak(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


def header_table(doc, left_lines, right_lines):
    tbl = doc.add_table(rows=1, cols=2)
    no_border(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    lc, rc = tbl.cell(0,0), tbl.cell(0,1)
    lc.width = Cm(8); rc.width = Cm(8)
    clr(lc); clr(rc)
    for txt, bd in left_lines:
        para = CP(lc, align=WD_ALIGN_PARAGRAPH.CENTER)
        R(para, txt, bold=bd)
    for txt, bd in right_lines:
        para = CP(rc, align=WD_ALIGN_PARAGRAPH.CENTER)
        R(para, txt, bold=bd)


# ══════════════════════════════════════════════════════════════════════════════
# PHẦN ĐỀ THI
# ══════════════════════════════════════════════════════════════════════════════

# LEFT_STANZAS / RIGHT_STANZAS được nạp tự động từ nho_ba.txt ở trên


def fill_poem_col(cell, stanzas):
    clr(cell)
    first = True
    for stanza in stanzas:
        if not first:
            blank = CP(cell, align=WD_ALIGN_PARAGRAPH.CENTER)
            blank.add_run('')
        first = False
        for line in stanza:
            para = CP(cell, align=WD_ALIGN_PARAGRAPH.CENTER)
            R(para, line, italic=True)


def build_exam(doc, cau2_vd, cau2_yc):
    header_table(doc,
        [('PHÒNG GD&ĐT ……………………', False),
         ('TRƯỜNG THCS ………………………', True)],
        [('KÌ THI CHỌN HỌC SINH GIỎI CẤP TRƯỜNG', True),
         ('NĂM HỌC 2025 – 2026', True),
         ('MÔN: NGỮ VĂN – LỚP 7', True)])
    P(doc, sb=2, sa=0)
    for txt in ['Thời gian làm bài: 90 phút, không kể thời gian giao đề',
                '(Đề thi có 01 trang)']:
        pr = P(doc, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=0)
        R(pr, txt, italic=True)
    P(doc, sb=0, sa=0)
    hrule(doc)

    # Câu 1
    pr1 = P(doc, sb=2, sa=1)
    R(pr1, 'Câu 1. ', bold=True); R(pr1, '(4,0 điểm) ', bold=True)
    R(pr1, 'Đọc bài thơ sau:')
    pt = P(doc, align=WD_ALIGN_PARAGRAPH.CENTER, sb=1, sa=1)
    R(pt, POEM_TITLE, bold=True)
    tp = doc.add_table(rows=1, cols=2)
    no_border(tp); tp.alignment = WD_TABLE_ALIGNMENT.CENTER
    tp.cell(0,0).width = Cm(7.5); tp.cell(0,1).width = Cm(7.5)
    fill_poem_col(tp.cell(0,0), LEFT_STANZAS)
    fill_poem_col(tp.cell(0,1), RIGHT_STANZAS)
    ps = P(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, sb=1, sa=2)
    R(ps, POEM_SOURCE, italic=True)
    pyc1 = P(doc, sb=0, sa=2, fi=1.0)
    R(pyc1, 'Viết đoạn văn khoảng '); R(pyc1, '300 chữ', bold=True)
    R(pyc1, ' ghi lại cảm xúc của em sau khi đọc bài thơ trên.')

    # Câu 2
    pr2 = P(doc, sb=2, sa=1)
    R(pr2, 'Câu 2. ', bold=True); R(pr2, '(6,0 điểm) ', bold=True)
    pvd = P(doc, sb=0, sa=1, fi=1.0)
    for txt, it in cau2_vd: R(pvd, txt, italic=it)
    pyc2 = P(doc, sb=0, sa=3, fi=1.0)
    for txt, bd, it in cau2_yc: R(pyc2, txt, bold=bd, italic=it)
    pend = P(doc, align=WD_ALIGN_PARAGRAPH.CENTER, sb=2, sa=0)
    R(pend, '--- Hết ---', bold=True)


# ══════════════════════════════════════════════════════════════════════════════
# PHẦN HƯỚNG DẪN CHẤM  –  format mới: 2 bảng riêng, cột Phần|Nội dung|Điểm
# ══════════════════════════════════════════════════════════════════════════════

# Chiều rộng 3 cột bảng thang điểm
CW = [Cm(3.2), Cm(10.0), Cm(1.8)]


def score_table(doc, rows_data):
    """
    Tạo bảng thang điểm.
    rows_data: list of (phan_txt, nd_items, diem_txt)
       phan_txt  : str  – cột Phần, ví dụ "a. Hình thức (0,5)"
       nd_items  : list[(text, bold, italic)] – nhiều đoạn trong cột Nội dung
       diem_txt  : str  – cột Điểm
    """
    tbl = doc.add_table(rows=len(rows_data)+1, cols=3)
    full_border(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, (w, txt) in enumerate(zip(CW, ['Phần', 'Nội dung cần đạt', 'Điểm'])):
        tbl.rows[0].cells[i].width = w
        clr(tbl.rows[0].cells[i])
        ph = CP(tbl.rows[0].cells[i], align=WD_ALIGN_PARAGRAPH.CENTER)
        R(ph, txt, bold=True)

    # Nội dung
    for ri, (phan, nd_items, diem) in enumerate(rows_data, 1):
        cells = tbl.rows[ri].cells
        for i, w in enumerate(CW): cells[i].width = w

        # Cột Phần
        clr(cells[0])
        pp = CP(cells[0], align=WD_ALIGN_PARAGRAPH.LEFT)
        R(pp, phan, bold=True)

        # Cột Nội dung cần đạt
        clr(cells[1])
        for txt, bd, it in nd_items:
            pr = CP(cells[1])
            R(pr, txt, bold=bd, italic=it)

        # Cột Điểm
        clr(cells[2])
        pd = CP(cells[2], align=WD_ALIGN_PARAGRAPH.CENTER)
        R(pd, diem, bold=bool(diem))


def luu_y(doc, lines):
    """Thêm đoạn Lưu ý sau bảng."""
    para = P(doc, sb=3, sa=1)
    R(para, 'Lưu ý: ', bold=True, italic=True)
    R(para, lines[0], italic=True)
    for line in lines[1:]:
        para2 = P(doc, sb=0, sa=1)
        R(para2, line, italic=True)


def ghi_chu_cho_diem(doc, rows):
    """Ghi chú cách cho điểm chi tiết Câu 2."""
    pg = P(doc, sb=8, sa=2)
    R(pg, 'Ghi chú cách cho điểm chi tiết Câu 2 (theo mức độ):', bold=True, underline=True)
    for prefix, rest in rows:
        pr = P(doc, sb=1, sa=1)
        R(pr, '- '); R(pr, prefix, bold=True); R(pr, rest)


# ──────────────────────────────────────────────────────────────────────────────
# NỘI DUNG CÂU 1 (CHUNG CHO CẢ 3 ĐỀ – bài thơ NHỚ BÀ)
# ──────────────────────────────────────────────────────────────────────────────

C1_ROWS = [
    ('a. Hình thức (0,5)',
     [('Đúng hình thức đoạn văn (mở – thân – kết), dung lượng khoảng 300 chữ.', False, False)],
     '0,5'),
    ('b. Mở đoạn (0,5)',
     [('Giới thiệu tác giả, tác phẩm; nêu ấn tượng chung về bài thơ.', False, False)],
     '0,5'),
    ('c. Thân đoạn (2,5)',
     [
      ('1. Cảm xúc về nội dung (1,5):', True, False),
      ('- Cảm nhận về hình ảnh người bà qua những chi tiết giản dị: bậu cửa, gầu giếng, '
       'chổi tre, bếp lửa… (0,5)', False, False),
      ('- Nỗi nhớ thương da diết khi bà đã khuất: nỗi đau, sự trống vắng, tiếc thương (0,5)', False, False),
      ('- Tình cảm gia đình sâu sắc, gắn bó với kỉ niệm tuổi thơ (0,5)', False, False),
      ('2. Cảm xúc về nghệ thuật (1,0):', True, False),
      ('- Thể thơ 5 chữ, nhịp điệu linh hoạt, giàu cảm xúc (0,25)', False, False),
      ('- Hình ảnh thơ bình dị, giàu sức gợi: "cơn gió thoảng về núi xanh", '
       '"vườn xưa xào xạc lá", "vàng mùa thu" (0,25)', False, False),
      ('- Biện pháp điệp ngữ "ngỡ bà" nhấn mạnh sự hòa quyện giữa bà với '
       'thiên nhiên, đất trời (0,25)', False, False),
      ('- Ngôn ngữ trong sáng, tha thiết, phù hợp với dòng cảm xúc hoài niệm (0,25)', False, False),
     ],
     '2,5'),
    ('d. Kết đoạn (0,25)',
     [('Khẳng định lại giá trị bài thơ, ý nghĩa với bản thân.', False, False)],
     '0,25'),
    ('e. Sáng tạo (0,25)',
     [('Diễn đạt mới mẻ, cảm xúc chân thành, có cách nhìn riêng.', False, False)],
     '0,25'),
]

C1_LUU_Y = [
    'Nếu đoạn văn viết quá ngắn (dưới 200 chữ) trừ 0,25; nếu viết lan man, '
    'thiếu trọng tâm trừ 0,5 – 1,0.',
]

GHI_CHU_CHUNG = [
    ('5,5 – 6,0: ', 'Đáp ứng đầy đủ yêu cầu, lập luận sắc sảo, dẫn chứng sinh động, sáng tạo tốt.'),
    ('4,5 – 5,25: ', 'Đáp ứng khá đầy đủ, lập luận rõ, dẫn chứng phù hợp, còn vài lỗi nhỏ.'),
    ('3,5 – 4,25: ', 'Đáp ứng cơ bản, bố cục rõ nhưng nội dung chưa sâu, dẫn chứng còn chung.'),
    ('2,5 – 3,25: ', 'Bài viết sơ sài, thiếu luận điểm, dẫn chứng nghèo nàn.'),
    ('Dưới 2,5: ',   'Không đạt yêu cầu, ý rời rạc, mắc nhiều lỗi.'),
]


# ──────────────────────────────────────────────────────────────────────────────
# DỮ LIỆU CÂU 2 RIÊNG TỪNG ĐỀ
# ──────────────────────────────────────────────────────────────────────────────

# ─── ĐỀ 1: Đọc sách ──────────────────────────────────────────────────────────
C2_ROWS_DE1 = [
    ('a. Bố cục (0,5)',
     [('Đầy đủ 3 phần: mở bài, thân bài, kết bài.', False, False)],
     '0,5'),
    ('b. Xác định vấn đề (0,25)',
     [('Đúng: vai trò và ý nghĩa của thói quen đọc sách đối với học sinh.', False, False)],
     '0,25'),
    ('c. Mở bài (0,5)',
     [('Dẫn dắt vấn đề; trích dẫn nhận định; nêu ý kiến tán thành của bản thân.', False, False)],
     '0,5'),
    ('d. Thân bài (4,0)',
     [
      ('1. Giải thích ý kiến (1,0):', True, False),
      ('- "Đọc sách" là tiếp thu kiến thức, giá trị cuộc sống từ sách (0,25)', False, False),
      ('- "Không chỉ để biết": đọc sách không đơn giản là ghi nhớ thông tin (0,25)', False, False),
      ('- "Còn để sống tốt hơn": đọc sách bồi dưỡng tâm hồn, nhân cách, kỹ năng sống (0,25)', False, False),
      ('- Ý kiến đề cao vai trò toàn diện, sâu sắc của văn hóa đọc (0,25)', False, False),
      ('2. Bàn luận về thói quen đọc sách (2,0):', True, False),
      ('- Đọc sách giúp mở rộng hiểu biết, tích lũy kiến thức, trau dồi vốn ngôn ngữ (0,5)', False, False),
      ('- Đọc sách bồi dưỡng tâm hồn: hiểu con người, biết cảm thông và chia sẻ (0,5)', False, False),
      ('- Đọc sách rèn luyện tư duy phản biện, kỹ năng tự học và sự tập trung (0,5)', False, False),
      ('- Dẫn chứng: danh nhân thành công nhờ đọc sách; học sinh tiêu biểu; liên hệ bản thân (0,5)', False, False),
      ('3. Bài học nhận thức và hành động (1,0):', True, False),
      ('- Xây dựng thói quen đọc sách mỗi ngày, chọn sách phù hợp lứa tuổi (0,5)', False, False),
      ('- Phê phán thái độ lười đọc sách, nghiện màn hình điện tử quá mức (0,5)', False, False),
     ],
     '4,0'),
    ('e. Kết bài (0,5)',
     [('Khẳng định lại ý nghĩa của việc đọc sách; liên hệ bản thân.', False, False)],
     '0,5'),
    ('f. Chính tả, ngữ pháp (0,25)',
     [('Đảm bảo đúng chính tả, dùng từ, đặt câu.', False, False)],
     '0,25'),
    ('g. Sáng tạo (0,5)',
     [('Lập luận độc đáo, dẫn chứng thuyết phục, có góc nhìn riêng phù hợp lứa tuổi.', False, False)],
     '0,5'),
]
C2_LUU_Y_DE1 = [
    'Nếu bài viết chỉ nêu chung chung, không có dẫn chứng hoặc dẫn chứng không phù hợp, trừ 0,5 – 1,0.',
    'Nếu viết quá sơ sài, mắc nhiều lỗi diễn đạt, không đạt yêu cầu về dung lượng, điểm tối đa không quá 3,0.',
]

# ─── ĐỀ 2: Tình bạn chân chính ───────────────────────────────────────────────
C2_ROWS_DE2 = [
    ('a. Bố cục (0,5)',
     [('Đầy đủ 3 phần: mở bài, thân bài, kết bài.', False, False)],
     '0,5'),
    ('b. Xác định vấn đề (0,25)',
     [('Đúng: tầm quan trọng của tình bạn chân chính đối với lứa tuổi học sinh.', False, False)],
     '0,25'),
    ('c. Mở bài (0,5)',
     [('Dẫn dắt vào vấn đề; trích dẫn nhận định; nêu ý kiến tán thành.', False, False)],
     '0,5'),
    ('d. Thân bài (4,0)',
     [
      ('1. Giải thích ý kiến (1,0):', True, False),
      ('- "Tình bạn chân chính": tình bạn thật sự, trong sáng, không vụ lợi (0,25)', False, False),
      ('- "Ngọn lửa sưởi ấm tâm hồn": tình bạn mang đến sự ấm áp, yêu thương, '
       'niềm vui tinh thần (0,25)', False, False),
      ('- "Điểm tựa giúp vượt qua thử thách": bạn bè là chỗ dựa vững chắc khi gặp '
       'khó khăn, thất bại (0,25)', False, False),
      ('- Ý kiến đề cao giá trị tinh thần và vai trò không thể thiếu của tình bạn '
       'chân chính (0,25)', False, False),
      ('2. Bàn luận về tầm quan trọng của tình bạn chân chính (2,0):', True, False),
      ('- Tình bạn chân chính mang lại động lực học tập, chia sẻ vui buồn, '
       'cùng nhau vượt qua khó khăn (0,5)', False, False),
      ('- Tình bạn tốt góp phần hình thành nhân cách lành mạnh, giúp tránh xa '
       'những cám dỗ tiêu cực (0,5)', False, False),
      ('- Tình bạn chân chính giúp ta trưởng thành, hiểu ý nghĩa của sự sẻ chia '
       'và lòng vị tha (0,5)', False, False),
      ('- Dẫn chứng: câu chuyện thực tế hoặc tác phẩm văn học về tình bạn; '
       'liên hệ bản thân (0,5)', False, False),
      ('3. Bài học nhận thức và hành động (1,0):', True, False),
      ('- Biết chọn bạn tốt; trân trọng và vun đắp tình bạn qua hành động '
       'cụ thể hằng ngày (0,5)', False, False),
      ('- Phê phán thái độ cơ hội, tình bạn giả tạo, vụ lợi trong '
       'cuộc sống (0,5)', False, False),
     ],
     '4,0'),
    ('e. Kết bài (0,5)',
     [('Khẳng định lại giá trị của tình bạn chân chính; liên hệ bản thân và ước mong.', False, False)],
     '0,5'),
    ('f. Chính tả, ngữ pháp (0,25)',
     [('Đảm bảo đúng chính tả, dùng từ, đặt câu.', False, False)],
     '0,25'),
    ('g. Sáng tạo (0,5)',
     [('Lập luận độc đáo, dẫn chứng thuyết phục, có góc nhìn riêng phù hợp lứa tuổi.', False, False)],
     '0,5'),
]
C2_LUU_Y_DE2 = [
    'Nếu bài viết chỉ nêu chung chung, không có dẫn chứng hoặc dẫn chứng không phù hợp, trừ 0,5 – 1,0.',
    'Nếu viết quá sơ sài, mắc nhiều lỗi diễn đạt, không đạt yêu cầu về dung lượng, điểm tối đa không quá 3,0.',
]

# ─── ĐỀ 3: Lòng biết ơn ──────────────────────────────────────────────────────
C2_ROWS_DE3 = [
    ('a. Bố cục (0,5)',
     [('Đầy đủ 3 phần: mở bài, thân bài, kết bài.', False, False)],
     '0,5'),
    ('b. Xác định vấn đề (0,25)',
     [('Đúng: ý nghĩa của lòng biết ơn trong cuộc sống học sinh.', False, False)],
     '0,25'),
    ('c. Mở bài (0,5)',
     [('Dẫn dắt vấn đề, trích dẫn câu nói, nêu ý kiến tán thành.', False, False)],
     '0,5'),
    ('d. Thân bài (4,0)',
     [
      ('1. Giải thích ý kiến (1,0):', True, False),
      ('- Lòng biết ơn là sự trân trọng, ghi nhớ và đền đáp công ơn (0,25)', False, False),
      ('- "Biến những gì ta có thành đủ": biết hài lòng với hiện tại, '
       'trân trọng những gì đang có (0,25)', False, False),
      ('- "Biến hỗn loạn thành trật tự": giúp ta sắp xếp cảm xúc, '
       'định hình giá trị sống (0,25)', False, False),
      ('- "Biến nhầm lẫn thành rõ ràng": phân biệt đúng – sai, '
       'nên – không nên trong ứng xử (0,25)', False, False),
      ('2. Bàn luận về ý nghĩa của lòng biết ơn với học sinh (2,0):', True, False),
      ('- Biết ơn cha mẹ: chăm ngoan, học tốt, chia sẻ việc nhà (0,5)', False, False),
      ('- Biết ơn thầy cô: kính trọng, lễ phép, nỗ lực học tập (0,5)', False, False),
      ('- Biết ơn bạn bè, những người giúp đỡ mình (0,25)', False, False),
      ('- Biết ơn cuộc sống: sống lạc quan, yêu đời, biết đủ (0,25)', False, False),
      ('- Dẫn chứng: gương học sinh vượt khó biết ơn cộng đồng; '
       'câu chuyện trong văn học, đời sống (0,5)', False, False),
      ('3. Bài học nhận thức và hành động (1,0):', True, False),
      ('- Lòng biết ơn cần thể hiện qua hành động cụ thể mỗi ngày (0,5)', False, False),
      ('- Phê phán thái độ vô ơn, thờ ơ, coi mọi điều là hiển nhiên (0,5)', False, False),
     ],
     '4,0'),
    ('e. Kết bài (0,5)',
     [('Khẳng định lại ý nghĩa của lòng biết ơn; liên hệ bản thân.', False, False)],
     '0,5'),
    ('f. Chính tả, ngữ pháp (0,25)',
     [('Đảm bảo đúng chính tả, dùng từ, đặt câu.', False, False)],
     '0,25'),
    ('g. Sáng tạo (0,5)',
     [('Lập luận độc đáo, dẫn chứng thuyết phục, có góc nhìn riêng phù hợp lứa tuổi.', False, False)],
     '0,5'),
]
C2_LUU_Y_DE3 = [
    'Nếu bài viết chỉ nêu chung chung, không có dẫn chứng hoặc dẫn chứng không phù hợp, trừ 0,5 – 1,0.',
    'Nếu viết quá sơ sài, mắc nhiều lỗi diễn đạt, không đạt yêu cầu về dung lượng, điểm tối đa không quá 3,0.',
]


# ══════════════════════════════════════════════════════════════════════════════
# HÀM BUILD HDC MỚI
# ══════════════════════════════════════════════════════════════════════════════

def build_hdc(doc, c2_vande_label, c2_rows, c2_luu_y):
    # Header
    header_table(doc,
        [('PHÒNG GD&ĐT ……………………', False),
         ('TRƯỜNG THCS ………………………', True)],
        [('HƯỚNG DẪN CHẤM ĐỀ THI HỌC SINH GIỎI', True),
         ('NĂM HỌC 2025 – 2026', True),
         ('MÔN: NGỮ VĂN – LỚP 7', True)])
    P(doc, sb=2, sa=0); hrule(doc)

    # I. Hướng dẫn chung
    ph = P(doc, sb=4, sa=2); R(ph, 'I. HƯỚNG DẪN CHUNG', bold=True)
    for i, txt in enumerate([
        'Giám khảo cần nắm vững yêu cầu của Hướng dẫn chấm để đánh giá tổng quát '
        'bài làm của học sinh, tránh cách chấm đếm ý cho điểm.',
        'Do đặc trưng của môn Ngữ văn, giám khảo cần chủ động, linh hoạt trong việc '
        'vận dụng đáp án và thang điểm; khuyến khích những bài viết có tư duy khoa học, '
        'sáng tạo, cảm xúc, có phong cách riêng.',
        'Điểm lẻ của toàn bài là 0,25 điểm.',
    ], 1):
        pr = P(doc, sb=1, sa=1); R(pr, f'{i}. {txt}')

    # II. Đáp án
    ph2 = P(doc, sb=6, sa=3); R(ph2, 'II. ĐÁP ÁN VÀ THANG ĐIỂM CHI TIẾT', bold=True)

    # ── Câu 1 ────────────────────────────────────────────────────────────────
    pc1 = P(doc, sb=2, sa=2)
    R(pc1, 'Câu 1. ', bold=True)
    R(pc1, '(4,0 điểm)', bold=True)
    score_table(doc, C1_ROWS)
    luu_y(doc, C1_LUU_Y)

    # ── Câu 2 ────────────────────────────────────────────────────────────────
    pc2 = P(doc, sb=6, sa=2)
    R(pc2, 'Câu 2. ', bold=True)
    R(pc2, '(6,0 điểm)', bold=True)
    score_table(doc, c2_rows)
    luu_y(doc, c2_luu_y)

    # Ghi chú cách cho điểm
    ghi_chu_cho_diem(doc, GHI_CHU_CHUNG)


# ══════════════════════════════════════════════════════════════════════════════
# TẠO FILE
# ══════════════════════════════════════════════════════════════════════════════

def make(path, cau2_vd, cau2_yc, c2_vande_label, c2_rows, c2_luu_y):
    doc = Document()
    setup(doc)
    build_exam(doc, cau2_vd, cau2_yc)
    pagebreak(doc)
    build_hdc(doc, c2_vande_label, c2_rows, c2_luu_y)
    doc.save(path)
    print(f'✔ {path}')


# ══════════════════════════════════════════════════════════════════════════════
# ĐỊNH NGHĨA CÁC ĐỀ
# ══════════════════════════════════════════════════════════════════════════════

ALL_EXAMS = [
  # ─── Đề 1: Đọc sách ──────────────────────────────────────────────────────
  dict(
    so_de=1,
    filename='ĐỀ HSG NGỮ VĂN 7 2025-2026 (Đề 1).doc',
    label='Đề 1 – Đọc sách',
    cau2_vd=[
        ('Có ý kiến cho rằng: ', False),
        ('"Đọc sách không chỉ là để biết mà còn là để sống tốt hơn."', True),
    ],
    cau2_yc=[
        ('Hãy viết bài văn khoảng ', False, False),
        ('400 đến 500 chữ', True, False),
        (' trình bày ', False, False),
        ('ý kiến tán thành', True, False),
        (' của em về ý kiến trên.', False, False),
    ],
    c2_vande_label='đọc sách',
    c2_rows=C2_ROWS_DE1,
    c2_luu_y=C2_LUU_Y_DE1,
  ),
  # ─── Đề 2: Tình bạn chân chính ───────────────────────────────────────────
  dict(
    so_de=2,
    filename='ĐỀ HSG NGỮ VĂN 7 2025-2026 (Đề 2).doc',
    label='Đề 2 – Tình bạn chân chính',
    cau2_vd=[
        ('Có ý kiến cho rằng: ', False),
        ('"Tình bạn chân chính là ngọn lửa sưởi ấm tâm hồn, là điểm tựa giúp chúng ta '
         'vượt qua mọi thử thách trong cuộc đời."', True),
    ],
    cau2_yc=[
        ('Hãy viết bài văn khoảng ', False, False),
        ('400 đến 500 chữ', True, False),
        (' trình bày ', False, False),
        ('ý kiến tán thành', True, False),
        (' của em và làm rõ ', False, False),
        ('tầm quan trọng của một tình bạn chân chính', False, True),
        (' đối với lứa tuổi học sinh.', False, False),
    ],
    c2_vande_label='tình bạn chân chính',
    c2_rows=C2_ROWS_DE2,
    c2_luu_y=C2_LUU_Y_DE2,
  ),
  # ─── Đề 3: Lòng biết ơn ──────────────────────────────────────────────────
  dict(
    so_de=3,
    filename='ĐỀ HSG NGỮ VĂN 7 2025-2026 (Đề 3).doc',
    label='Đề 3 – Lòng biết ơn',
    cau2_vd=[
        ('Nhà văn Melody Beattie từng viết: ', False),
        ('"Lòng biết ơn biến những gì ta có thành đủ; biến sự hỗn loạn thành trật tự; '
         'biến sự nhầm lẫn thành sự rõ ràng."', True),
    ],
    cau2_yc=[
        ('Hãy viết bài văn khoảng ', False, False),
        ('400 đến 500 chữ', True, False),
        (' trình bày ', False, False),
        ('ý kiến tán thành', True, False),
        (' của em về câu nói trên, từ đó bàn về ', False, False),
        ('ý nghĩa của lòng biết ơn trong cuộc sống của học sinh', False, True),
        ('.', False, False),
    ],
    c2_vande_label='lòng biết ơn',
    c2_rows=C2_ROWS_DE3,
    c2_luu_y=C2_LUU_Y_DE3,
  ),
]


# ══════════════════════════════════════════════════════════════════════════════
# ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════════

def run(exams_to_run: list[int] | None = None, validate_only: bool = False):
    """Validate điểm và/hoặc tạo file."""
    from validate_scores import validate_all

    # Kiểm tra điểm
    ok = validate_all(C1_ROWS, [
        (e['label'], e['c2_rows']) for e in ALL_EXAMS
    ])
    if not ok:
        print('\n⛔ Dừng lại! Sửa lỗi điểm trước khi tạo file.')
        return
    if validate_only:
        return

    # Tạo thư mục output
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Lọc đề cần tạo
    selected = ALL_EXAMS if not exams_to_run else [
        e for e in ALL_EXAMS if e['so_de'] in exams_to_run
    ]
    if not selected:
        print(f'⚠ Không tìm thấy đề số: {exams_to_run}')
        return

    for exam in selected:
        out_path = OUTPUT_DIR / exam['filename']
        make(
            str(out_path),
            cau2_vd       = exam['cau2_vd'],
            cau2_yc       = exam['cau2_yc'],
            c2_vande_label= exam['c2_vande_label'],
            c2_rows       = exam['c2_rows'],
            c2_luu_y      = exam['c2_luu_y'],
        )
    print(f'\n✅ Hoàn thành! File đã lưu vào: {OUTPUT_DIR}')


if __name__ == '__main__':
    ap = argparse.ArgumentParser(
        description='Tạo bộ đề HSG Ngữ Văn 7 (2025-2026)'
    )
    ap.add_argument(
        '--de', nargs='+', type=int, metavar='N',
        help='Số đề cần tạo (1, 2, 3). Bỏ trống = tạo tất cả.'
    )
    ap.add_argument(
        '--validate', action='store_true',
        help='Chỉ kiểm tra điểm, không tạo file Word.'
    )
    args = ap.parse_args()
    run(exams_to_run=args.de, validate_only=args.validate)
